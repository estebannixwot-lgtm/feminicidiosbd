from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Titulo
title = document.add_heading('Plataforma de Análisis de Datos Geográficos de Feminicidios en Colombia', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Info
document.add_paragraph('Informe Metodológico y Técnico\nAutor: Esteban Pinto\nInstitución: Universidad\nFecha: 27 de Abril de 2026').alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_heading('Resumen', level=1)
document.add_paragraph('El presente documento expone la metodología, el diseño y la implementación de una plataforma integral de análisis de datos espaciales enfocada en la problemática del feminicidio en Colombia. Este proyecto responde a la necesidad de visualizar y analizar de forma estructurada los crecientes casos de violencia de género en el territorio nacional, facilitando la toma de decisiones informadas y la formulación de políticas públicas.')

document.add_heading('1. Introducción', level=1)
document.add_paragraph('En Colombia, los casos de feminicidio han presentado un preocupante incremento, demandando herramientas tecnológicas eficaces para su seguimiento y análisis. La presente iniciativa plantea la construcción de una base de datos geoespacial y un aplicativo web (tablero de comando) que permita la consulta dinámica de la distribución geográfica y estadística de estos incidentes. El objetivo principal es estructurar la información disponible, proveniente de entidades como el Observatorio de Feminicidios Colombia, en un modelo lógico robusto soportado por tecnologías de vanguardia como PostgreSQL, PostGIS, Express.js y Leaflet.')

document.add_heading('2. Requerimientos y Diseño Conceptual', level=1)
document.add_heading('2.1 Modelo de Requerimientos de Datos', level=2)
document.add_paragraph('Se identificaron las siguientes entidades principales requeridas para la comprensión del fenómeno:', style='List Bullet')
document.add_paragraph('Víctima y Agresor: Información demográfica, edad, nivel educativo, relación previa e incidentes o antecedentes.', style='List Bullet')
document.add_paragraph('Caso de Feminicidio: Atributos espaciotemporales (fecha, ubicación), tipo de feminicidio y tipificación legal inicial.', style='List Bullet')
document.add_paragraph('Entidades Espaciales: División política administrativa de Colombia (Departamentos y Municipios).', style='List Bullet')
document.add_paragraph('Estadísticas Nacionales y Departamentales: Resúmenes anuales y tendencias de la violencia.', style='List Bullet')

document.add_heading('2.2 Modelo Conceptual y Lógico', level=2)
document.add_paragraph('El modelo conceptual y el modelo lógico (entregados en formato PDF) definen las relaciones de integridad referencial:')
document.add_paragraph('La tabla CASO_FEMINICIDIO actúa como entidad central (Tabla de Hechos), enlazándose con las dimensiones de VICTIMA, AGRESOR, MUNICIPIO y FUENTE_INFORMACION.', style='List Bullet')
document.add_paragraph('Se diseñó un esquema relacional normalizado con soporte espacial, donde las geometrías (geom) de departamentos, municipios y casos específicos se almacenan bajo los estándares del Open Geospatial Consortium (OGC) usando la extensión PostGIS.', style='List Bullet')

document.add_heading('3. Metodología de Implementación', level=1)
document.add_heading('3.1 Tecnologías Utilizadas', level=2)
document.add_paragraph('Base de Datos: PostgreSQL 16.12 con la extensión PostGIS habilitada, configurada en el puerto 5433.', style='List Bullet')
document.add_paragraph('Backend: Node.js y Express, con el conector pg para la API REST.', style='List Bullet')
document.add_paragraph('Frontend: HTML5, CSS3, y JavaScript Vanilla. Diseño Glassmorphism con Chart.js y Leaflet.js.', style='List Bullet')

document.add_heading('3.2 Construcción y Poblamiento de la Base de Datos', level=2)
document.add_paragraph('Debido a la incompatibilidad de versiones en la cabecera de la copia de seguridad recibida (FEMINICIDIO_COLOMBIA.backup), se procedió a:')
document.add_paragraph('Reconstrucción del Esquema (DDL): Se elaboró y ejecutó un script SQL completo que plasma las relaciones del modelo lógico.', style='List Number')
document.add_paragraph('Extracción Documental: Se extrajo información del Observatorio de Feminicidios Colombia (Ej. 621 casos reportados en 2025).', style='List Number')
document.add_paragraph('Poblamiento de Datos: Los datos se cargaron en ESTADISTICAS_NACIONALES y ESTADISTICAS_FEMINICIDIOS_POR_DEPARTAMENTO.', style='List Number')

document.add_heading('4. Resultados y Aplicativo Web (Tablero de Comando)', level=1)
document.add_paragraph('El sistema resultante presenta:')
document.add_paragraph('Panel lateral (Sidebar): Muestra resúmenes estadísticos.', style='List Bullet')
document.add_paragraph('Gráficos en tiempo real: Diagramas de barras interactivos (Chart.js).', style='List Bullet')
document.add_paragraph('Mapa Temático: Mapa base interactivo con Leaflet, cuyo radio de círculos es proporcional al número de feminicidios.', style='List Bullet')

document.add_heading('5. Conclusiones', level=1)
document.add_paragraph('El desarrollo del proyecto logró estructurar de manera coherente la base de datos espacial solicitada, manteniendo estricto apego al modelo lógico aportado. La reconstrucción mediante DDL y su integración web asegura la viabilidad del proyecto, sirviendo como herramienta vital para el análisis de la problemática.')

document.add_heading('Referencias', level=1)
document.add_paragraph('Observatorio de Feminicidios Colombia. (2025). Informes estadísticos de violencia de género en el marco colombiano.')
document.add_paragraph('OSGeo. (2023). PostGIS Spatial Extender for PostgreSQL.')

document.save('c:/U 2026/ChatBotICG/PROY_BD/ENTREGABLES/Informe_Proyecto_Feminicidios_APA.docx')
